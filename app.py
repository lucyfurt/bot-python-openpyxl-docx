from openpyxl import load_workbook
from docx import Document
from datetime import datetime

# Carregar a planilha Excel
planilha_aulas = load_workbook('./plano.xlsx')
pagina_aulas = planilha_aulas['Sheet1'] 

# Iterar pelas linhas da planilha
for linha in pagina_aulas.iter_rows(min_row=2, values_only=True):
    try:
        # Desempacotar apenas os valores necessários
        periodo, curso, nivel, turno, componente, ch_semanal, ch_total, obj_geral, ementa, obj_especificos, conteudos, metodologia, recursos, avaliacao, instrumentos, criterios_avaliacao, recuperacao, final, projetos, referencias, adaptacao, obs, dia, mes, *extras = linha
    except ValueError:
        print("Erro ao desempacotar linha:", linha)
        continue
    
    # Criar o documento do Word
    arquivo_word = Document()
    arquivo_word.add_heading('Plano de Ensino', 0)

    # Adicionar o texto ao documento
    texto_plano = f"""
    PLANO DE ENSINO - Curso Técnico

    1. Identificação
        1.1 Docente: Luciana Nathalia Morais Furtado
        1.2 Período Letivo: {periodo}
        1.3 Curso:  {curso}
        1.4 Nível: {nivel}
        1.5 Turma/turno: {turno}
        1.6 Componente curricular: {componente}   | Carga horária semanal: {ch_semanal}   Carga horária total: {ch_total}

    2. Objetivo (s) Geral (is):
       {obj_geral} 

    3. Ementa:
       {ementa}  

    4. Objetivos Específicos:
        {obj_especificos}

    5. Conteúdos:
        {conteudos}

    6. Metodologia:
        {metodologia}

    7. Recursos:
        {recursos}

    8. Avaliação:
        {avaliacao}
        
        8.1 Instrumentos avaliativos a serem usados pelo (a) docente:
            {instrumentos}
        8.2 Critérios de avaliação:
            {criterios_avaliacao}
            
        8.3 Recuperação Paralela e final:
        {recuperacao}
        {final}

    9. Projetos e/ou visitas técnicas:
        {projetos}

    10. Referências:
        {referencias}

    11. Adaptação às necessidades específicas:
        {adaptacao}

    12. Observações gerais:
        {obs}

    Santa Inês, {dia},{mes}, 2024	



    Assinado eletronicamente
    ___________________________
    Docente

    ___________________________
    Setor pedagógico

    _____________________________ 
    Coordenação do curso
    """

    arquivo_word.add_paragraph(texto_plano)

    # Salvar o documento
    arquivo_word.save(f'./planos/plano_{componente}.docx')
